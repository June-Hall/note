import os
import re
from datetime import datetime
import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CSS = """
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+SC&display=swap');
body { font-family: 'Noto Sans SC', 'Microsoft YaHei', sans-serif; font-size: 12pt; line-height: 1.8; margin: 2cm; color: #222; }
h1 { font-size: 20pt; border-bottom: 2px solid #333; padding-bottom: 6px; }
h2 { font-size: 15pt; color: #1a5276; margin-top: 20px; }
h3 { font-size: 13pt; color: #154360; }
blockquote { background: #f5f5f5; border-left: 4px solid #2196F3; padding: 8px 12px; margin: 8px 0; }
strong { color: #c0392b; }
code { background: #f0f0f0; padding: 2px 5px; border-radius: 3px; }
pre { background: #f0f0f0; padding: 12px; border-radius: 4px; }
"""

def to_markdown(notes_md: str, output_path: str):
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(notes_md)

def to_pdf(notes_md: str, output_path: str):
    """PDF 导出需要 GTK，如果未安装会抛出异常"""
    try:
        from weasyprint import HTML
        body = markdown.markdown(notes_md, extensions=["extra", "toc"])
        html = f"<html><head><meta charset='utf-8'><style>{CSS}</style></head><body>{body}</body></html>"
        HTML(string=html).write_pdf(output_path)
    except Exception as e:
        # 保存 HTML 作为替代方案
        html_path = output_path.replace('.pdf', '.html')
        body = markdown.markdown(notes_md, extensions=["extra", "toc"])
        html = f"<html><head><meta charset='utf-8'><style>{CSS}</style></head><body>{body}</body></html>"
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html)
        raise Exception(f"PDF 生成失败（需要 GTK），已生成 HTML: {html_path}")

def _add_formatted_runs(paragraph, text):
    """解析行内 **加粗**、==高亮==、<span class="highlight-xxx"> 标记，写入 Word run。"""
    # 先把前端的高亮 span 转换为 ==文本== 形式，并记录颜色
    # <span class="highlight-red">x</span> -> 用占位标记
    text = re.sub(r'<span class="highlight-red">(.*?)</span>', r'⟦RED⟧\1⟦/RED⟧', text)
    text = re.sub(r'<span class="highlight-yellow">(.*?)</span>', r'⟦YEL⟧\1⟦/YEL⟧', text)
    text = re.sub(r'<span class="highlight-green">(.*?)</span>', r'⟦GRN⟧\1⟦/GRN⟧', text)
    text = re.sub(r'<span class="highlight-blue">(.*?)</span>', r'⟦BLU⟧\1⟦/BLU⟧', text)
    text = re.sub(r'==(.+?)==', r'⟦YEL⟧\1⟦/YEL⟧', text)

    # 分词：加粗与高亮标记
    pattern = re.compile(
        r'(\*\*.+?\*\*|⟦RED⟧.*?⟦/RED⟧|⟦YEL⟧.*?⟦/YEL⟧|⟦GRN⟧.*?⟦/GRN⟧|⟦BLU⟧.*?⟦/BLU⟧)'
    )
    parts = pattern.split(text)

    color_map = {
        'RED': (RGBColor(0xC0, 0x39, 0x2B), WD_COLOR_INDEX.RED),
        'YEL': (None, WD_COLOR_INDEX.YELLOW),
        'GRN': (None, WD_COLOR_INDEX.BRIGHT_GREEN),
        'BLU': (RGBColor(0x19, 0x76, 0xD2), WD_COLOR_INDEX.TURQUOISE),
    }

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            matched = False
            for tag, (font_color, highlight) in color_map.items():
                if part.startswith(f'⟦{tag}⟧'):
                    inner = part[len(tag) + 2:-(len(tag) + 3)]
                    run = paragraph.add_run(inner)
                    if font_color:
                        run.font.color.rgb = font_color
                    run.font.highlight_color = highlight
                    matched = True
                    break
            if not matched:
                paragraph.add_run(part)


def _shade_cell(cell, hex_color: str):
    """给表格单元格设置背景色（用于高亮框）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _add_box(doc, title: str, lines: list[str], fill: str, accent: RGBColor, icon: str = ""):
    """添加一个带背景色的高亮框（单格表格实现）。"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _shade_cell(cell, fill)

    # 标题行
    head = cell.paragraphs[0]
    head_run = head.add_run(f"{icon} {title}".strip())
    head_run.bold = True
    head_run.font.size = Pt(12)
    head_run.font.color.rgb = accent

    # 内容行
    for line in lines:
        line = line.strip()
        if not line:
            continue
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        _add_formatted_runs(p, line.lstrip("•-* "))

    doc.add_paragraph()  # 框后留白
    return table


def to_docx(notes_md: str, output_path: str, subject: str = "课程笔记"):
    doc = Document()

    # 页边距
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 正文默认字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ===== 封面标题 =====
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    trun = title_p.add_run(f"《{subject}》课堂笔记")
    trun.bold = True
    trun.font.size = Pt(24)
    trun.font.color.rgb = RGBColor(0x1F, 0x24, 0x21)

    date_p = doc.add_paragraph()
    drun = date_p.add_run(f"整理时间：{datetime.now().strftime('%Y-%m-%d')}")
    drun.font.size = Pt(9)
    drun.font.color.rgb = RGBColor(0x8A, 0x8A, 0x80)

    # 分隔线
    _add_hr(doc)

    # ===== 正文解析 =====
    lines = notes_md.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 教师重点块：以「教师重点」标题开头，收集后续列表项
        if re.match(r'^#{1,4}\s*.*教师重点', stripped) or stripped in ("教师重点", "**教师重点**"):
            block, i = _collect_block(lines, i + 1)
            _add_box(doc, "教师重点", block, fill="FCE8D8", accent=RGBColor(0xD9, 0x6B, 0x2B), icon="★")
            continue

        # 我的笔记块
        if re.match(r'^#{1,4}\s*.*我的笔记', stripped) or stripped in ("我的笔记", "**我的笔记**"):
            block, i = _collect_block(lines, i + 1)
            _add_box(doc, "我的笔记", block, fill="E3EEF9", accent=RGBColor(0x19, 0x76, 0xD2), icon="✎")
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            h = doc.add_heading("", level=2)
            r = h.add_run("◉ " + line[3:])
            r.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.style = doc.styles["Quote"] if "Quote" in doc.styles else p.style
            _add_formatted_runs(p, line[2:])
        elif re.match(r'^\s*[-*•]\s+', line):
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, re.sub(r'^\s*[-*•]\s+', '', line))
        elif stripped:
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)
        i += 1

    # ===== 页脚 =====
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "NotesAI 自动生成"
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x8A, 0x8A, 0x80)

    doc.save(output_path)


def _collect_block(lines, start):
    """从 start 开始收集连续的列表/正文行，遇到下一个标题或空段落停止。"""
    block = []
    i = start
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith("#") or re.match(r'^#{1,4}\s', s):
            break
        if s == "" and block:
            # 单个空行视为块结束
            break
        if s:
            block.append(s)
        i += 1
    return block, i


def _add_hr(doc):
    """添加一条水平分隔线。"""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E2D9C8')
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
